import os
import re
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_document_formatting(doc):
    # Set paper size to Quarto (8.0 in x 10.0 in)
    # Set margins: Left = 1.5 inches, Top, Bottom, Right = 1.0 inch
    section = doc.sections[0]
    section.page_width = Inches(8.0)
    section.page_height = Inches(10.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.left_margin = Inches(1.5)
    
    # Configure default Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

def add_h1(doc, text, is_first=False, prevent_page_break=False):
    # Insert a page break before any major heading unless it is the very first heading
    # or if we want to prevent it (e.g. sequential headings like CHAPTER ONE / INTRODUCTION)
    if not is_first and not prevent_page_break:
        doc.add_page_break()
        
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_format = p.paragraph_format
    # Preliminary headings are single spaced if TOC/Abstract, otherwise double. Let's make all main headers centered double spaced.
    p_format.line_spacing = 2.0
    p_format.space_before = Pt(18)
    p_format.space_after = Pt(12)
    p_format.first_line_indent = Inches(0)
    p_format.keep_with_next = True
    
    run = p.add_run(text.upper())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    return p

def add_h2_h3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_format = p.paragraph_format
    p_format.line_spacing = 2.0
    p_format.space_before = Pt(12)
    p_format.space_after = Pt(6)
    p_format.first_line_indent = Inches(0)
    p_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    return p

def add_paragraph_with_inline_formatting(doc, text, style='Normal', alignment=None, is_list_item=False, is_heading=False, single_space=False, is_title_page=False):
    if is_list_item:
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph()
        
    if alignment is not None:
        p.alignment = alignment
    elif is_title_page:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    p_format = p.paragraph_format
    p_format.line_spacing = 1.0 if (single_space and not is_title_page) else 2.0
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    
    # First-line indent: only for normal body paragraphs (not list items, not headings, not centered placeholders, not TOC dotted lines, not Title Page)
    is_special_line = (
        is_list_item or 
        is_heading or 
        is_title_page or
        alignment == WD_ALIGN_PARAGRAPH.CENTER or 
        "..." in text or 
        "BY" in text or 
        "MATRIC" in text or 
        "DEGREE IN" in text or
        "__________________" in text
    )
    if not is_special_line:
        p_format.first_line_indent = Inches(0.5)
    else:
        p_format.first_line_indent = Inches(0)
        
    # Split text by bold markers **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            run = p.add_run(bold_text)
            run.bold = True
        else:
            # Check for italic markers
            subparts = re.split(r'(\*.*?\*|__.*?__|_.*?_)', part)
            for subpart in subparts:
                if (subpart.startswith('*') and subpart.endswith('*')) or (subpart.startswith('_') and subpart.endswith('_')):
                    italic_text = subpart[1:-1]
                    run = p.add_run(italic_text)
                    run.italic = True
                else:
                    p.add_run(subpart)
                    
    # Format all runs to Times New Roman (size 16 bold if title page, size 12 unless heading)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        if is_title_page:
            run.font.size = Pt(16)
            run.bold = True
        elif not is_heading:
            run.font.size = Pt(12)
    return p

def build_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Format headers
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0  # Tables are single-spaced
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.first_line_indent = Inches(0)
        for run in p.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
    # Format rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < len(row_cells):
                cell_val = cell_text.strip()
                row_cells[c_idx].text = cell_val
                p = row_cells[c_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.0  # Tables are single-spaced
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.first_line_indent = Inches(0)
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

def compile_markdown_to_docx(md_path, docx_path):
    print(f"Compiling {md_path} to {docx_path}...")
    doc = docx.Document()
    set_document_formatting(doc)
    
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_table = False
    table_headers = []
    table_rows = []
    
    # Track current page context to toggle double vs single spacing
    # Title Page should be double spaced per the guidelines, so we exclude it from single-spaced sections.
    single_spaced_sections = {
        "CERTIFICATION", "TABLE OF CONTENTS", "LIST OF TABLES", 
        "LIST OF FIGURES", "ABBREVIATIONS AND ACRONYMS", "ABSTRACT"
    }
    current_section = ""
    is_first_h1 = True
    just_added_h1 = False
    
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        
        # Handle Table parsing
        if line.startswith("|"):
            in_table = True
            cells = [c.strip() for c in line.split("|")[1:-1]]
            
            # Check if this is separator line (e.g. |:---|:---|)
            if all(re.match(r'^:?-+:?$', c) for c in cells):
                idx += 1
                continue
                
            if not table_headers:
                table_headers = cells
            else:
                table_rows.append(cells)
                
            idx += 1
            continue
        else:
            if in_table:
                # Flush the table
                if table_headers:
                    build_table(doc, table_headers, table_rows)
                    # Add spacer paragraph after table
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.0 if current_section in single_spaced_sections else 2.0
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                table_headers = []
                table_rows = []
                in_table = False
            
        if not line:
            idx += 1
            continue
            
        # Parse Headings
        if line.startswith("# "):
            heading_text = line[2:].strip()
            current_section = heading_text.upper()
            prevent_page_break = just_added_h1
            add_h1(doc, heading_text, is_first=is_first_h1, prevent_page_break=prevent_page_break)
            is_first_h1 = False
            just_added_h1 = True
        elif line.startswith("## "):
            heading_text = line[3:]
            add_h2_h3(doc, heading_text)
            just_added_h1 = False
        elif line.startswith("### "):
            heading_text = line[4:]
            add_h2_h3(doc, heading_text)
            just_added_h1 = False
        elif line.startswith("#### "):
            heading_text = line[5:]
            add_h2_h3(doc, heading_text)
            just_added_h1 = False
            
        # Parse Bullet Lists
        elif line.startswith("- ") or line.startswith("* "):
            list_text = line[2:]
            single_space = current_section in single_spaced_sections
            add_paragraph_with_inline_formatting(doc, list_text, is_list_item=True, single_space=single_space)
            just_added_h1 = False
            
        # Parse Screenshot Placeholders (rendered in italics or centered bold/italic for clarity)
        elif line.startswith("[INSERT SCREENSHOT:"):
            single_space = current_section in single_spaced_sections
            add_paragraph_with_inline_formatting(doc, line, alignment=WD_ALIGN_PARAGRAPH.CENTER, single_space=single_space)
            just_added_h1 = False
            
        # Parse normal paragraph
        else:
            single_space = current_section in single_spaced_sections
            # Check for centered alignments in Title page
            alignment = None
            is_tp = (current_section == "TITLE PAGE")
            if is_tp or "UNIVERSITY OF ILORIN" in line or "FACULTY OF COMMUNICATION" in line:
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_paragraph_with_inline_formatting(doc, line, alignment=alignment, single_space=single_space, is_title_page=is_tp)
            just_added_h1 = False
            
        idx += 1
        
    # Final flush in case document ends with a table
    if in_table and table_headers:
        build_table(doc, table_headers, table_rows)
        
    try:
        doc.save(docx_path)
        print(f"Successfully compiled: {docx_path}")
    except PermissionError:
        print(f"\n[!] ERROR: Permission denied when saving to {docx_path}.")
        print("    Please close this document in Microsoft Word and run compile_chapters.py again.\n")

def main():
    writeup_dir = r"c:\wamp64\www\finalyearproject\writeup"
    
    # We will compile the chapters
    if os.path.exists(os.path.join(writeup_dir, "Chapter_One_AEP.md")):
        compile_markdown_to_docx(
            os.path.join(writeup_dir, "Chapter_One_AEP.md"),
            os.path.join(writeup_dir, "Chapter_One_AEP.docx")
        )
    
    if os.path.exists(os.path.join(writeup_dir, "Chapter_Four.md")):
        compile_markdown_to_docx(
            os.path.join(writeup_dir, "Chapter_Four.md"),
            os.path.join(writeup_dir, "Chapter_Four.docx")
        )
    
    if os.path.exists(os.path.join(writeup_dir, "Chapter_Five.md")):
        compile_markdown_to_docx(
            os.path.join(writeup_dir, "Chapter_Five.md"),
            os.path.join(writeup_dir, "Chapter_Five.docx")
        )
        
    # Compile IDS chapters
    if os.path.exists(os.path.join(writeup_dir, "Chapter_Four_IDS.md")):
        compile_markdown_to_docx(
            os.path.join(writeup_dir, "Chapter_Four_IDS.md"),
            os.path.join(writeup_dir, "Chapter_Four_IDS.docx")
        )
        
    if os.path.exists(os.path.join(writeup_dir, "Chapter_Five_IDS.md")):
        compile_markdown_to_docx(
            os.path.join(writeup_dir, "Chapter_Five_IDS.md"),
            os.path.join(writeup_dir, "Chapter_Five_IDS.docx")
        )
        
    if os.path.exists(os.path.join(writeup_dir, "Chapter_One_IDS_Project.md")):
        compile_markdown_to_docx(
            os.path.join(writeup_dir, "Chapter_One_IDS_Project.md"),
            os.path.join(writeup_dir, "Chapter_One_IDS_Project.docx")
        )

if __name__ == "__main__":
    main()
