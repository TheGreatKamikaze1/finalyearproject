import os
import docx
import pypdf

def extract_docx(file_path):
    print(f"Extracting DOCX: {file_path}")
    doc = docx.Document(file_path)
    fullText = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        fullText.append(para.text)
        
    # Extract tables
    for table in doc.tables:
        fullText.append("\n--- Table ---")
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                # Avoid duplicate text if cells are merged
                txt = cell.text.strip().replace('\n', ' ')
                row_text.append(txt)
            fullText.append(" | ".join(row_text))
        fullText.append("-------------\n")
        
    return '\n'.join(fullText)

def extract_pdf(file_path):
    print(f"Extracting PDF: {file_path}")
    reader = pypdf.PdfReader(file_path)
    fullText = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ''
        fullText.append(f"\n--- Page {i+1} ---")
        fullText.append(text)
    return '\n'.join(fullText)

def main():
    writeup_dir = r"c:\wamp64\www\finalyearproject\writeup"
    out_dir = os.path.join(writeup_dir, "extracted_text")
    os.makedirs(out_dir, exist_ok=True)
    
    files = {
        "Chapter_One_AEP.docx": "Chapter_One_AEP.txt",
        "Final Year Project Template_Updated.docx": "Final_Year_Project_Template_Updated.txt",
        "FYP Temitope (1) up.docx": "FYP_Temitope_1_up.txt",
        "furthered salawu oluwatobi dorcas.pdf": "furthered_salawu_oluwatobi_dorcas.txt"
    }
    
    for filename, out_filename in files.items():
        file_path = os.path.join(writeup_dir, filename)
        out_path = os.path.join(out_dir, out_filename)
        
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            continue
            
        try:
            if filename.endswith(".docx"):
                text = extract_docx(file_path)
            elif filename.endswith(".pdf"):
                text = extract_pdf(file_path)
            else:
                continue
                
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(text)
            print(f"Successfully saved to: {out_path}")
        except Exception as e:
            print(f"Error processing {filename}: {e}")

if __name__ == "__main__":
    main()
